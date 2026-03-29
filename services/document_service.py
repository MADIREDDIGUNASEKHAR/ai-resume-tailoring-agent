"""
document_service.py
-------------------
Responsible for:
  - Reading the candidate DOCX resume → plain text
  - Converting tailored LLM text → a well-formatted DOCX
  - Saving the output file with a descriptive name
"""
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import settings
from models.job_model import Job
from utils.logger import get_logger
from utils.helpers import ensure_dir, sanitize_filename

logger = get_logger(__name__)

# Section heading keywords used to detect logical resume sections
_SECTION_KEYWORDS = (
    "summary", "objective", "skills", "experience", "employment",
    "work history", "projects", "education", "certifications",
    "awards", "publications", "languages", "interests",
)


def read_resume(path: str | None = None) -> str:
    """
    Extract the full text from the candidate's DOCX resume.
    Returns a single string with paragraphs separated by newlines.
    Raises FileNotFoundError if the file is missing.
    """
    path = path or settings.RESUME_FILE
    if not os.path.exists(path):
        raise FileNotFoundError(f"Resume file not found: {path}")

    logger.info(f"Reading resume from: {path}")
    doc = Document(path)
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also grab text from tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in lines:
                    lines.append(cell_text)

    resume_text = "\n".join(lines)
    logger.info(f"Extracted {len(lines)} paragraph(s) / {len(resume_text)} chars from resume.")
    return resume_text


def _is_section_heading(line: str) -> bool:
    """Heuristic: is this line a resume section heading?"""
    stripped = line.strip().lower().rstrip(":")
    return stripped in _SECTION_KEYWORDS or (
        len(line) < 60
        and line.strip().endswith(":")
        and any(kw in stripped for kw in _SECTION_KEYWORDS)
    )


def _is_name_line(line: str, position: int) -> bool:
    """First non-blank line is treated as the candidate's name."""
    return position == 0 and len(line.split()) <= 5


def save_resume(tailored_text: str, job: Job) -> str:
    """
    Convert the tailored plain-text resume to a formatted DOCX file.
    Returns the absolute path to the saved file.
    """
    ensure_dir(settings.OUTPUT_DIR)

    filename = sanitize_filename(f"resume_{job.safe_filename_part()}.docx")
    output_path = os.path.join(settings.OUTPUT_DIR, filename)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = [ln for ln in tailored_text.splitlines()]  # preserve all, filter below

    non_empty_idx = 0  # counts non-blank lines to identify name line
    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            # Blank spacer paragraph
            doc.add_paragraph()
            continue

        if _is_name_line(line, non_empty_idx):
            # Candidate name – large, bold, centred
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            non_empty_idx += 1
            continue

        if non_empty_idx in (1, 2, 3) and not _is_section_heading(line):
            # Contact / location line(s) under the name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            non_empty_idx += 1
            continue

        if _is_section_heading(line):
            # Section heading – bold, blue, with top spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xBF)
            non_empty_idx += 1
            continue

        # Bullet points (lines starting with -, *, •)
        if re.match(r"^[-*•]\s+", line):
            content = re.sub(r"^[-*•]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(content)
            run.font.size = Pt(11)
            non_empty_idx += 1
            continue

        # Regular body text
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(11)
        non_empty_idx += 1

    doc.save(output_path)
    logger.info(f"Saved tailored resume → {output_path}")
    return output_path
